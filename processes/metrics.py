import os
from datetime import date, datetime

from django.conf import settings
from django.http import HttpResponse
from django.urls import reverse
from docx import Document
from assisteds.models import AssistedDocumentModel, AssistedsModel
from cooperators.models import CooperatorsModel
from processes import utils
from processes.models import ProcessesModel, ProcessDocumentModel, ProcessesNotesModel
from processes.forms import ProcessesDocumentForm,  ProcessesNotesForm
from requireds.forms import RequiredsForm
from responsibles.models import ResponsiblesDocumentsModel


def get_context_data(request, instance, context):
    document_assisted = AssistedDocumentModel.objects.filter(id_assisted=instance.id_assisted)
    document_responsible = ResponsiblesDocumentsModel.objects.filter(id_responsible=instance.id_responsible)
    required = ProcessesModel.objects.filter(id_required=instance.id_required).first()
    note_waring = ProcessesNotesModel.objects.filter(id_process=instance.pk, status='0')


    if request.POST:
        context['requireds_form'] = RequiredsForm(request.POST)
        context['form_notes'] = ProcessesNotesForm(request.POST)
    else:
        context['requireds_form'] = RequiredsForm(initial={'id_assisted': instance.id_assisted})
        context['form_notes'] = ProcessesNotesForm(initial={'id_process': instance.pk})

    action_mapping = utils.action_mapping()
    context['categories'] = action_mapping
    context['required'] = required
    context['document_assisted'] = document_assisted
    context['document_responsibles'] = document_responsible
    context['note_waring'] = note_waring
    return context

def get_context_data_document(instance, context):
    documents_assisted = AssistedDocumentModel.objects.filter(id_assisted=instance.get('id_assisted') )
    documents_responsible = ResponsiblesDocumentsModel.objects.filter(id_responsible=instance.get('id_responsible'))
    context['documents_assisted'] = documents_assisted
    context['documents_responsible'] = documents_responsible
    return context

from datetime import datetime
import os
from django.http import HttpResponse
from django.conf import settings
from docx import Document

def generate_document(processes):
    # Obter dados do processo, assistido e advogado
    process = ProcessesModel.objects.filter(protocol=processes.protocol).first()
    assisted = AssistedsModel.objects.filter(id=processes.id_assisted_id).first()
    advocate = CooperatorsModel.objects.filter(id=processes.id_advocate_id).first()

    if not (process and assisted and advocate):
        print("Erro: Processo, assistido ou advogado não encontrados.")
        return HttpResponse("Dados não encontrados", status=404)

    # Carregar o arquivo do modelo
    file_path = os.path.join(settings.MEDIA_ROOT, 'archives', 'procuracao.docx')
    if not os.path.exists(file_path):
        print("Erro: Arquivo procuracao.docx não encontrado no caminho especificado.")
        return HttpResponse("Arquivo de modelo não encontrado", status=404)

    doc = Document(file_path)
    day = datetime.now().day
    year = datetime.now().year
    meses = [
        "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
        "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
    ]
    month = meses[datetime.now().month - 1]

    # Substituições dos dados
    substitutions = {
        "{{name_assisted}}": str(assisted.name).upper(),
        "{{nacionalidade_assisted}}": "brasileiro(a)".upper(),
        "{{status_assisted}}": str(assisted.status_civil).upper(),
        "{{working_assisted}}": str(assisted.employee).upper(),
        "{{cpf_assisted}}": str(assisted.cpf).upper(),
        "{{rg_assisted}}": str(assisted.rg).upper(),
        "{{expedition_assisted}}": str(assisted.organ).upper(),
        "{{address_assisted}}": str(assisted.address).upper(),
        "{{district_assisted}}": str(assisted.district).upper(),
        "{{city_assisted}}": str(assisted.city).upper(),
        "{{uf_assisted}}": str(assisted.uf).upper(),
        "{{cep_assisted}}": "48.925-000".upper(),
        "{{name_advocate}}": str(advocate.name).upper(),
        "{{nacionalidade_advocate}}": "brasileiro(a)".upper(),
        "{{status_advocate}}": str(advocate.status_civil).upper(),
        "{{number_order_advocate}}": str(advocate.number_order).upper(),
        "{{address_advocate}}": str(advocate.address).upper(),
        "{{district_advocate}}": str(advocate.district).upper(),
        "{{city_advocate}}": str(advocate.city).upper(),
        "{{uf_advocate}}": str(advocate.uf).upper(),
        "{{cep_advocate}}": "48.925-000".upper(),
        "{{email_advocate}}": str(advocate.email).upper(),
        "{{day}}": str(day),
        "{{month}}": month,
        "{{year}}": str(year)
    }

    # Função para substituir texto em runs preservando formatação
    def replace_text_in_paragraph(paragraph):
        # Concatenar todo o texto das runs
        full_text = "".join(run.text for run in paragraph.runs)

        # Substituir os placeholders no texto concatenado
        for placeholder, replacement in substitutions.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, replacement)

        # Limpar as runs e recriar com o texto atualizado
        if full_text:
            paragraph.clear()
            paragraph.add_run(full_text)

    # Substituir variáveis em parágrafos e células de tabelas
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)

    # Salvar o arquivo preenchido
    output_path = os.path.join(settings.MEDIA_ROOT, 'archives', 'procuracao_preenchido.docx')
    doc.save(output_path)

    # Retornar o arquivo como resposta HTTP
    with open(output_path, 'rb') as file:
        response = HttpResponse(file.read(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="procuracao_preenchido.docx"'
        return response


def generate_document_hipo(processes):
    # Obter dados do processo, assistido e advogado
    process = ProcessesModel.objects.filter(protocol=processes.protocol).first()
    assisted = AssistedsModel.objects.filter(id=processes.id_assisted_id).first()
    advocate = CooperatorsModel.objects.filter(id=processes.id_advocate_id).first()

    if not (process and assisted and advocate):
        print("Erro: Processo, assistido ou advogado não encontrados.")
        return HttpResponse("Dados não encontrados", status=404)

    # Carregar o arquivo do modelo
    file_path = os.path.join(settings.MEDIA_ROOT, 'archives', 'hiposuficiencia.docx')
    if not os.path.exists(file_path):
        print("Erro: Arquivo procuracao.docx não encontrado no caminho especificado.")
        return HttpResponse("Arquivo de modelo não encontrado", status=404)

    doc = Document(file_path)
    day = datetime.now().day
    year = datetime.now().year
    meses = [
        "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
        "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
    ]
    month = meses[datetime.now().month - 1]

    # Substituições dos dados
    substitutions = {
        "{{name_assisted}}": str(assisted.name).upper(),
        "{{nacionalidade_assisted}}": "brasileiro(a)".upper(),
        "{{status_assisted}}": str(assisted.status_civil).upper(),
        "{{working_assisted}}": str(assisted.employee).upper(),
        "{{cpf_assisted}}": str(assisted.cpf).upper(),
        "{{rg_assisted}}": str(assisted.rg).upper(),
        "{{expedition_assisted}}": str(assisted.organ).upper(),
        "{{address_assisted}}": str(assisted.address).upper(),
        "{{district_assisted}}": str(assisted.district).upper(),
        "{{city_assisted}}": str(assisted.city).upper(),
        "{{uf_assisted}}": str(assisted.uf).upper(),
        "{{cep_assisted}}": "48.925-000".upper(),
        "{{day}}": str(day),
        "{{month}}": month,
        "{{year}}": str(year)
    }

    # Função para substituir texto em runs preservando formatação
    def replace_text_in_paragraph(paragraph):
        # Concatenar todo o texto das runs
        full_text = "".join(run.text for run in paragraph.runs)

        # Substituir os placeholders no texto concatenado
        for placeholder, replacement in substitutions.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, replacement)

        # Limpar as runs e recriar com o texto atualizado
        if full_text:
            paragraph.clear()
            paragraph.add_run(full_text)

    # Substituir variáveis em parágrafos e células de tabelas
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)

    # Salvar o arquivo preenchido
    output_path = os.path.join(settings.MEDIA_ROOT, 'archives', 'hiposuficiencia_preenchido.docx')
    doc.save(output_path)

    # Retornar o arquivo como resposta HTTP
    with open(output_path, 'rb') as file:
        response = HttpResponse(file.read(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="hiposuficiencia_preenchido.docx"'
        return response




