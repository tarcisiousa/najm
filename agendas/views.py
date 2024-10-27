import json
import os
from datetime import date, datetime

from django.contrib.auth.mixins import LoginRequiredMixin
from docx import Document
from django.conf import settings
from django.utils import timezone
from django.urls import reverse
from django.db.models import Count, Sum
from django.db.models.functions import TruncDate
from django.http import JsonResponse, HttpResponseRedirect, HttpResponse
from django.shortcuts import render, redirect
from django.urls import reverse
from django.views.generic import ListView, CreateView, DetailView, UpdateView, DeleteView
from processes.models import ProcessesModel
from .models import AgendaModel
from .forms import AgendaForm
from .models import AgendaModel
from django.http import JsonResponse
from django.db.models import Count
from .models import AgendaModel
from assisteds.models import AssistedsModel

class AgendaList(LoginRequiredMixin, ListView):
    template_name = 'agendas_list.html'
    model = AgendaModel
    context_object_name = 'agendas'

    def get_queryset(self):
        date_now = date.today()
        agendas = AgendaModel.objects.filter(date_at=date_now)
        search_date = self.request.GET.get('search-date', '').replace('/', '-')

        if search_date:
            agendas = AgendaModel.objects.filter(date_at=search_date)
        print(len(agendas))
        return agendas

    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, *kwargs)
        agendas_not_served = AgendaModel.objects.filter(
            status='0',
            date_at__lt=timezone.now()  # substitua 'data' pelo nome do campo de data no seu modelo
        )

        context['not_serveds'] = agendas_not_served

        return context

    def post(self, *args, **kwargs):
        protocol = self.request.POST.get('protocol', '')
        current_protocol = AgendaModel.objects.filter(protocol=protocol).first()

        if protocol:
            try:
                current_assisted = AssistedsModel.objects.get(pk=current_protocol.id_assisted.pk)
            except AssistedsModel.DoesNotExist:
                return HttpResponse("Assistido não encontrado", status=404)

            # Certifique-se de que está atribuindo a instância, não um QuerySet
            new_agenda = AgendaModel(protocol=protocol, id_assisted=current_assisted, status='0')
            new_agenda.save()

            return HttpResponseRedirect(reverse('agendas_update', kwargs={'pk': new_agenda.pk}))

        return HttpResponse("Protocolo não encontrado", status=404)


class AgendaCreate(LoginRequiredMixin, CreateView):
    model = AgendaModel
    form_class = AgendaForm
    template_name = 'agendas_create.html'
    success_url = '/agendas/list/'

    def form_valid(self, form):
        return super().form_valid(form)


class AgendaDetail(LoginRequiredMixin, DetailView):
    model = AgendaModel
    template_name = 'agendas_detail.html'

    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        agenda = self.get_object()
        process = ProcessesModel.objects.filter(protocol=agenda.protocol)

        context['process'] = process

        return context


class AgendaDetailProtocol(LoginRequiredMixin, DetailView):
    model = AgendaModel
    template_name = 'agendas_detail_protocol.html'

    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        agenda = self.get_object()
        process = ProcessesModel.objects.filter(protocol=agenda.protocol)

        context['process'] = process

        return context


    def post(self, *args, **kwargs):
        agenda_current = self.get_object()
        process = ProcessesModel.objects.filter(protocol=agenda_current.protocol).first()
        file_path = os.path.join(settings.MEDIA_ROOT, 'archives', 'ficha_atendimento.docx')
        doc = Document(file_path)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{{protocol}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{{protocol}}", str(process.protocol).upper())
                        if "{{assisted}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{{assisted}}", str(process.id_assisted).upper())
                        if "{{responsible}}" in paragraph.text:
                            if process.id_responsible != None:
                                id_responsible = process.id_responsible
                            else:
                                id_responsible = 'Desnecessário'
                            paragraph.text = paragraph.text.replace("{{responsible}}", str(id_responsible).upper())
                        if "{{date}}" in paragraph.text:
                            date_at = date.strftime(agenda_current.date_at, '%d/%m/%Y')
                            paragraph.text = paragraph.text.replace("{{date}}", str(date_at).upper())
                        if "{{time}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("{{time}}", f'{agenda_current.time_at.upper()}hr')
                        if "{{number_process}}" in paragraph.text:
                            print(process.number_process)
                            if process.number_process != None:
                                number_process = str(process.number_process).upper()
                            else:
                                number_process = ''
                                print(number_process)
                            paragraph.text = paragraph.text.replace("{{number_process}}", number_process)

        # Salvar o arquivo novo
        output_path = os.path.join(settings.MEDIA_ROOT, 'archives', 'ficha_atendimento_preenchido.docx')
        doc.save(output_path)

        # Criar resposta HTTP para download do arquivo
        with open(output_path, 'rb') as file:
            response = HttpResponse(file.read(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="ficha_atendimento_preenchido.docx"'
            return response

class AgendaUpdate(LoginRequiredMixin, UpdateView):
    model = AgendaModel
    form_class = AgendaForm
    template_name = 'agendas_update.html'

    def post(self, request, *args, **kwargs):
        print('inicou')
        self.object = self.get_object()  # Carrega o objeto antes de qualquer operação
        form = self.get_form()

        if form.is_valid():
            print('form valido')
            # Pegando os dados diretamente do POST
            id_assisted = request.POST.get('id_assisted')
            protocol = request.POST.get('protocol')
            date_selected = request.POST.get('data')
            time_selected = request.POST.get('time-selected')

            agenda = form.save(commit=False)
            agenda.id_assisted_id = id_assisted
            agenda.protocol = protocol
            agenda.date_at = date_selected
            agenda.time_at = datetime.now().strftime('%H:%M') if time_selected is None else time_selected
            agenda.status = '0'
            agenda.save()
            
            #add date_ate is service in processes
            processes = ProcessesModel.objects.filter(protocol=protocol).first()
            processes.date_at = date_selected
            processes.save()

            return self.form_valid(form)
        else:
            print(form.errors)
            # Se o formulário não for válido, retorne os erros
            return self.form_invalid(form)

    def get_success_url(self):
        # Use o nome da URL definida em urls.py, que é 'agendas_detail_protocol'
        return reverse('agendas_detail_protocol', kwargs={'pk': self.object.pk})


class AgendaDelete(LoginRequiredMixin, DeleteView):
    model = AgendaModel
    template_name = 'agendas_delete.html'
    success_url = '/agendas/list/'

    def delete(self, request, *args, **kwargs):
        agenda = self.get_object()
        try:
            process = ProcessesModel.objects.get(protocolo=agenda.protocolo)
            process.delete()
        except ProcessesModel.DoesNotExist:
            print("Nenhum processo encontrado para o agendamento.")

        return super().delete(request, *args, **kwargs)


def horary_vagos_view(LoginRequiredMixin, request):
    data_str = request.GET.get('data')

    if data_str:
        # Converter a string de data para um objeto de data
        data = datetime.strptime(data_str, '%Y-%m-%d').date()

        # Buscar todas as agendas do dia
        agendas_do_dia = AgendaModel.objects.filter(date_at=data)

        # Lista de horários totais disponíveis
        time_total = ['08:00', '08:30', '09:00', '09:30', '10:00', '10:30', '11:00', '11:30', '12:00', '12:30']

        # Pegar os horários já ocupados no dia
        time_used = [agenda.time_at for agenda in agendas_do_dia]

        # Encontrar os horários livres
        time_free = [hora for hora in time_total if hora not in time_used]

        return JsonResponse(time_free, safe=False)

    return JsonResponse([], safe=False)

def eventos_view(LoginRequiredMixin, request):
    # Obter a data atual
    today = date.today()

    # Pegar todas as agendas
    agendas = AgendaModel.objects.all()

    # Agrupar por date_at e contar as ocorrências de cada data
    contagem_por_data = AgendaModel.objects.values('date_at').annotate(total=Count('date_at')).order_by('date_at')
    days = []
    service = []

    # Lista de horários totais disponíveis no dia
    time_total = ['08:00', '08:30', '09:00', '09:30', '10:00', '10:30', '11:00', '11:30', '12:00', '12:30']

    # Exibir os resultados
    for data in contagem_por_data:
        days.append(data['date_at'])
        service.append(data['total'])
        time_used = []
        for item in agendas:
            if item.date_at == data['date_at']:
                time_used.append(item.time_at)

    def check(horary_ocupados, data_evento):
        # Verificar se data_evento é None
        if data_evento is None:
            return 'Data inválida'

        # Se a data já passou, retornar 'Fechado' em vermelho
        if data_evento < today:
            return '<span style="color:white;">Fechado</span>'

        if len(horary_ocupados) == len(time_total):  # Se todos os horários estão ocupados
            return 'Fechado'
        elif len(time_total) == 0:
            return 'Horários vagos:<br>Todos os Horários:'
        else:
            # Identificar os horários livres
            time_free_result = [hora for hora in time_total if hora not in horary_ocupados]
            # Juntar os horários livres com ' - ' e definir a quebra de linha manual após 3 horários
            time_free = ' - '.join(time_free_result[:3]) + '<br>' + ' - '.join(time_free_result[3:])
            return f'Horários vagos:<br>{time_free}'

    # Filtrar dias válidos, evitando None
    valid_days = [day for day in days if day is not None]

    # Criar eventos com a descrição dos horários ou a mensagem "Fechado"
    events = [
        {
            'title': '',
            'start': str(day),
            'description': check([item.time_at for item in agendas if item.date_at == day], day),
        }
        for day in valid_days
    ]

    return JsonResponse(events, safe=False)

def preencher_documento(LoginRequiredMixin, data):
    # Carregar o arquivo existente
    file_path = os.path.join(settings.MEDIA_ROOT, 'arquivos', 'seu_arquivo.docx')
    doc = Document(file_path)

    # Substituir marcadores pelos dados do banco
    for paragraph in doc.paragraphs:
        if "{{protocolo}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{protocolo}}", data['protocolo'])
        if "{{numero_processo}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{numero_processo}}", data['numero_processo'])
        if "{{date}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{date}}", data['data_primeiro_atendimento'])
        if "{{time}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{time}}", data['data_primeiro_atendimento'])
        # Continue para outros campos...

    # Salvar o arquivo novo
    output_path = os.path.join(settings.MEDIA_ROOT, 'arquivos', 'documento_preenchido.docx')
    doc.save(output_path)

    return output_path


    